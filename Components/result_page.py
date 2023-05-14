from PyQt6.QtCore import Qt
from PyQt6.QtWidgets import QMainWindow, QTableWidgetItem
from GUI.result_window_ui import Ui_MainWindow
from db.Database import Database
from docx import Document
from Components.message_dialog import MessageDialog
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side


class ResultPage(QMainWindow):
    def __init__(self, username, users, MainWindow, data, sub_types_txt, title='Заголовок'):
        super().__init__()
        self.title = title
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.database = Database()
        self.username = username
        self.users = users
        self.MainWindow = MainWindow
        self.ui.name.setText(username)
        self.ui.menu_back.clicked.connect(self.menu_back)
        self.MessageDialog = None
        self.ui.result_on_word.clicked.connect(self.result_on_word)
        self.ui.result_on_excel.clicked.connect(self.result_on_excel)
        self.sub_types_txt = sub_types_txt
        self.data = data
        self.init_table()

    def menu_back(self):
        self.close()
        self.MainWindow.show()

    def init_table(self):
        self.ui.res_table.setRowCount(len(self.data))
        self.ui.res_table.setColumnCount(len(self.data[0]))
        self.ui.res_table.setHorizontalHeaderLabels(self.sub_types_txt)
        self.ui.title.setText(self.title)
        for i, row in enumerate(self.data):
            for j, cell in enumerate(row):
                item = QTableWidgetItem(str(cell))
                item.setFlags(
                    Qt.ItemFlag.ItemIsSelectable | Qt.ItemFlag.ItemIsEditable | Qt.ItemFlag.ItemIsEnabled | Qt.ItemFlag.ItemIsDragEnabled)
                self.ui.res_table.setItem(i, j, item)
        self.ui.res_table.resizeColumnsToContents()
        self.ui.res_table.resizeRowsToContents()

    def result_on_word(self):
        document = Document()
        # Устанавливаем ориентацию альбомной ориентации
        section = document.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        # Задаем заголовок таблицы
        title = self.title
        # Добавляем заголовок перед таблицей
        document.add_heading(title, 0)
        # Создаем таблицу
        table = document.add_table(rows=len(self.data) + 1, cols=len(self.sub_types_txt),
                                   style='Table Grid')
        # Заполняем первую строку заголовками таблицы
        for j in range(len(self.sub_types_txt)):
            cell = table.cell(0, j)
            cell.text = self.sub_types_txt[j]
            cell.paragraphs[0].runs[0].bold = True  # жирный шрифт для заголовков
        # Заполняем таблицу
        for i in range(len(self.data)):
            for j in range(len(self.sub_types_txt)):
                cell = table.cell(i + 1, j)
                cell.text = str(self.data[i][j])
        # Сохраняем документ
        now = datetime.now()
        date_string = now.strftime("%d.%m.%Y")
        fileName = title + '_' + date_string + '.docx'
        document.save('genered_reports/' + fileName)
        self.MessageDialog = MessageDialog('Создан документ с именем ' + fileName)
        self.MessageDialog.show()

    def result_on_excel(self):
        # Создаем новый документ
        workbook = Workbook()
        # Добавляем активный лист
        worksheet = workbook.active
        # Задаем заголовок таблицы
        title = self.title
        # Задаем значения заголовков
        headers = self.sub_types_txt
        # Задаем данные таблицы
        data_table = self.data

        # Заполняем заголовки таблицы
        for column, header in enumerate(headers, start=1):
            cell = worksheet.cell(row=2, column=column, value=header)
            cell.font = Font(bold=True)
            alignment = Alignment(horizontal='center', vertical='center')
            cell.alignment = alignment
            cell.border = Border(left=Side(border_style='thin', color='000000'),
                                 right=Side(border_style='thin', color='000000'),
                                 top=Side(border_style='thin', color='000000'),
                                 bottom=Side(border_style='thin', color='000000'))

        # Заполняем таблицу данными
        for row, data in enumerate(data_table, start=1):
            for column, value in enumerate(data, start=1):
                cell = worksheet.cell(row=row + 2, column=column, value=value)
                alignment = Alignment(horizontal='center', vertical='center')
                cell.alignment = alignment
                cell.border = Border(left=Side(border_style='thin', color='000000'),
                                     right=Side(border_style='thin', color='000000'),
                                     top=Side(border_style='thin', color='000000'),
                                     bottom=Side(border_style='thin', color='000000'))
        for column_cells in worksheet.columns:
            # Вычисляем максимальный размер ячейки в столбце
            length = max(len(str(cell.value)) for cell in column_cells)
            # Устанавливаем ширину столбца
            worksheet.column_dimensions[column_cells[0].column_letter].width = length + 2
        # Вставляем заголовок перед таблицей
        worksheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
        worksheet.cell(row=1, column=1, value=title).alignment = Alignment(horizontal='center', vertical='center')
        # Сохраняем документ
        now = datetime.now()
        date_string = now.strftime("%d.%m.%Y")
        fileName = title + '_' + date_string + '.xlsx'
        workbook.save('genered_reports/' + fileName)
        self.MessageDialog = MessageDialog('Создан документ с именем ' + fileName)
        self.MessageDialog.show()
